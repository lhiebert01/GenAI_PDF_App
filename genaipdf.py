import warnings
import sys
import os
import base64
from PyPDF2 import PdfReader
import google.generativeai as genai
from langchain_community.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from datetime import datetime
import streamlit as st
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Suppress LangChain deprecation warnings
warnings.filterwarnings('ignore', category=DeprecationWarning)

# Set page config at the very beginning
st.set_page_config(
    page_title="Gemini Pro Advanced PDF Analysis & Chat",
    page_icon="📚",
    layout="wide"
)

# Load environment variables and configure API
load_dotenv(override=True)
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

def get_pdf_text(pdf_docs):
    """Extract text from multiple PDF documents"""
    text = ""
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

def get_text_chunks(text):
    """Split text into manageable chunks"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=10000,
        chunk_overlap=1000
    )
    chunks = text_splitter.split_text(text)
    return chunks

def get_vector_store(text_chunks):
    """Create vector store from text chunks"""
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")
    return vector_store

def process_pdfs(pdf_docs):
    """Process uploaded PDF files and create vector store"""
    raw_text = get_pdf_text(pdf_docs)
    text_chunks = get_text_chunks(raw_text)
    vector_store = get_vector_store(text_chunks)
    return vector_store

def create_qa_chain(vector_store):
    """Create an enhanced Question-Answering chain"""
    improved_prompt = PromptTemplate(
        input_variables=["context", "question"],
        template="""You are an advanced AI analyst tasked with providing comprehensive, well-structured responses. Follow these guidelines:

1. ANALYSIS APPROACH:
   - Thoroughly analyze all provided context
   - Synthesize information into clear, logical sections
   - Present information in easily digestible formats (tables, lists, categories)
   - Make connections between related concepts
   - Highlight key insights and patterns

2. RESPONSE STRUCTURE:
   - Start with a brief executive summary
   - Organize information into clear sections with headers
   - Use tables when comparing or categorizing information
   - Create detailed, hierarchical lists when appropriate
   - Include relevant examples and explanations

3. DEPTH OF COVERAGE:
   - Provide extensive detail for each important point
   - Explore multiple angles and perspectives
   - Connect concepts to broader themes
   - Identify implications and applications

Context from documents:
{context}

Question: {question}

Provide a comprehensive response that:
1. Answers the question using ONLY information from the context
2. Organizes the information in a clear, structured format
3. Uses tables and lists where appropriate
4. Highlights key insights and patterns
5. Makes connections between related concepts

If certain information is not available in the context, clearly state this and provide relevant insights from general knowledge in a separate section marked "Additional Insights (Beyond Context)"."""
    )
    
    llm = ChatGoogleGenerativeAI(
        model="gemini-1.5-flash-latest",
        temperature=0.4,
        max_output_tokens=8192,
        top_p=0.9,
        top_k=40
    )
    
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=vector_store.as_retriever(search_kwargs={"k": 10}),
        chain_type_kwargs={"prompt": improved_prompt}
    )
    
    return qa_chain

def get_binary_file_downloader_html(bin_file, file_label='File'):
    """Generate a download link for a file"""
    with open(bin_file, 'rb') as f:
        data = f.read()
    bin_str = base64.b64encode(data).decode()
    href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{os.path.basename(bin_file)}">Download {file_label}</a>'
    return href

def save_chat_history_txt(chat_history):
    """Save chat history to a text file"""
    if not chat_history:
        return None
        
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"chat_history_{timestamp}.txt"
    
    try:
        with open(filename, "w", encoding="utf-8") as f:
            for interaction in chat_history:
                f.write(f"Question:\n{interaction['question']}\n\n")
                f.write(f"Response:\n{interaction['response']}\n")
                f.write("\n" + "="*50 + "\n\n")
        return filename
    except Exception as e:
        raise Exception(f"Error saving text file: {str(e)}")

def save_chat_history_docx(chat_history):
    """Save chat history to a Word document"""
    if not chat_history:
        return None
        
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"chat_history_{timestamp}.docx"
    
    try:
        doc = docx.Document()
        
        title = doc.add_heading('Gemini Pro Advanced PDF Analysis Chat History', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for interaction in chat_history:
            question_heading = doc.add_heading('Question:', level=1)
            question_para = doc.add_paragraph(interaction['question'])
            
            response_heading = doc.add_heading('Response:', level=1)
            response_para = doc.add_paragraph(interaction['response'])
            
            doc.add_paragraph('=' * 50)
        
        doc.save(filename)
        return filename
    except Exception as e:
        raise Exception(f"Error saving Word document: {str(e)}")

def save_chat_history_docx_formatted(chat_history):
    """Save chat history to a Word document with proper formatting conversion from Markdown"""
    if not chat_history:
        return None
        
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"chat_history_formatted_{timestamp}.docx"
    
    try:
        doc = docx.Document()
        
        # Add title
        title = doc.add_heading('Gemini Pro Advanced PDF Analysis Chat History', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for interaction in chat_history:
            # Add question section
            question_heading = doc.add_heading('Question:', level=1)
            question_para = doc.add_paragraph(interaction['question'])
            
            # Add response section
            response_heading = doc.add_heading('Response:', level=1)
            
            # Split content into sections for processing
            sections = interaction['response'].split('\n\n')
            
            for section in sections:
                if '|' in section and '-|-' in section:  # Table detection
                    # Process table
                    rows = [row.strip() for row in section.split('\n') if row.strip() and '|-' not in row]
                    if rows:
                        # Count columns
                        header_cells = [cell.strip() for cell in rows[0].split('|') if cell.strip()]
                        num_cols = len(header_cells)
                        
                        # Create table
                        table = doc.add_table(rows=1, cols=num_cols)
                        table.style = 'Table Grid'
                        
                        # Add header row
                        header_row = table.rows[0].cells
                        for idx, cell_text in enumerate(header_cells):
                            header_row[idx].text = cell_text.strip('*')  # Remove markdown bold
                            paragraph = header_row[idx].paragraphs[0]
                            run = paragraph.runs[0]
                            run.bold = True
                        
                        # Add data rows
                        for row_text in rows[1:]:
                            cells = [cell.strip() for cell in row_text.split('|') if cell.strip()]
                            row_cells = table.add_row().cells
                            for idx, cell_text in enumerate(cells):
                                # Handle bold text in cells
                                text = cell_text.strip('*')  # Remove markdown bold
                                row_cells[idx].text = text
                                if cell_text.startswith('*') and cell_text.endswith('*'):
                                    row_cells[idx].paragraphs[0].runs[0].bold = True
                else:
                    # Process non-table content
                    lines = section.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        
                        # Check for section headers (bold text at start of line)
                        if line.startswith('**') and line.endswith('**'):
                            text = line.strip('*')
                            heading = doc.add_heading(text, level=2)
                        elif line.startswith('·') or line.startswith('* ') or line.startswith('- '):
                            # Process bullet points
                            # Remove bullet point marker and leading/trailing spaces
                            text = line.lstrip('·').lstrip('*').lstrip('-').strip()
                            
                            # Create bullet point paragraph
                            p = doc.add_paragraph(style='List Bullet')
                            
                            # Handle bold text within bullet points
                            if text.startswith('**') and text.endswith('**'):
                                # Bold text without the markdown
                                text = text.strip('*')
                                run = p.add_run(text)
                                run.bold = True
                            else:
                                # Process mixed formatting within bullet points
                                parts = text.split('**')
                                for i, part in enumerate(parts):
                                    if part:  # Skip empty parts
                                        run = p.add_run(part)
                                        run.bold = (i % 2 == 1)  # Bold for odd-indexed parts
                        else:
                            # Regular paragraph with possible mixed formatting
                            p = doc.add_paragraph()
                            # Split by bold markers
                            parts = line.split('**')
                            for i, part in enumerate(parts):
                                if part:  # Skip empty parts
                                    run = p.add_run(part)
                                    run.bold = (i % 2 == 1)  # Bold for odd-indexed parts
            
            # Add separator
            doc.add_paragraph('_' * 50)
        
        doc.save(filename)
        return filename
    except Exception as e:
        raise Exception(f"Error saving formatted Word document: {str(e)}")
    
def main():
    """Main Streamlit application"""
    # Initialize session state
    if 'vector_store' not in st.session_state:
        st.session_state.vector_store = None
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    
    st.title("📚 Gemini Pro Advanced PDF Analysis Companion  ✧ ✧ 📕📗📘📙💻🧑‍💻🌍 ✧ ✧ ")
    
    # Sidebar for document upload and processing
    with st.sidebar:
        st.header("📤 Document Upload")
        pdf_docs = st.file_uploader("Upload PDFs", type=['pdf'], accept_multiple_files=True)
        
        if st.button("Process Documents"):
            if not pdf_docs:
                st.warning("Please upload PDF documents first.")
            else:
                with st.spinner("Processing documents..."):
                    try:
                        st.session_state.vector_store = process_pdfs(pdf_docs)
                        st.success("✅ Documents processed successfully!")
                    except Exception as e:
                        st.error(f"Error: {str(e)}")
        
        st.header("💾 Export Chat History")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("Save as TXT"):
                try:
                    if st.session_state.chat_history:
                        with st.spinner("Creating text file..."):
                            save_path = save_chat_history_txt(st.session_state.chat_history)
                            if save_path:
                                st.markdown(
                                    get_binary_file_downloader_html(save_path, 'Text File'), 
                                    unsafe_allow_html=True
                                )
                                os.remove(save_path)
                    else:
                        st.warning("No chat history to save")
                except Exception as e:
                    st.error(f"Error saving file: {str(e)}")
        
        with col2:
            if st.button("Save as DOCX"):
                try:
                    if st.session_state.chat_history:
                        with st.spinner("Creating Word document..."):
                            save_path = save_chat_history_docx(st.session_state.chat_history)
                            if save_path:
                                st.markdown(
                                    get_binary_file_downloader_html(save_path, 'Word Document'), 
                                    unsafe_allow_html=True
                                )
                                os.remove(save_path)
                    else:
                        st.warning("No chat history to save")
                except Exception as e:
                    st.error(f"Error saving file: {str(e)}")
        
        with col3:
            if st.button("Save as Formatted DOCX"):
                try:
                    if st.session_state.chat_history:
                        with st.spinner("Creating formatted Word document..."):
                            save_path = save_chat_history_docx_formatted(st.session_state.chat_history)
                            if save_path:
                                st.markdown(
                                    get_binary_file_downloader_html(save_path, 'Formatted Word Document'), 
                                    unsafe_allow_html=True
                                )
                                os.remove(save_path)
                    else:
                        st.warning("No chat history to save")
                except Exception as e:
                    st.error(f"Error saving file: {str(e)}")
    
    # Main chat interface
    st.header("🤖 World's Most Advanced Intelligent PDF Document Analysis 🌍 ✧ ✧")
    
    with st.form(key='question_form'):
        user_question = st.text_area("Enter your question:", height=100)
        submit_button = st.form_submit_button("Submit")
        
        if submit_button and user_question:
            if not st.session_state.vector_store:
                st.warning("⚠️ Please process documents first")
            else:
                try:
                    qa_chain = create_qa_chain(st.session_state.vector_store)
                    with st.spinner("Generating response..."):
                        response = qa_chain.invoke({"query": user_question})
                        st.session_state.chat_history.insert(0, {
                            "question": user_question,
                            "response": response['result']
                        })
                except Exception as e:
                    st.error(f"Error: {str(e)}")
    
    # Display chat history
    if st.session_state.chat_history:
        st.header("💬 Conversation History")
        for idx, interaction in enumerate(st.session_state.chat_history):
            with st.expander(f"Q&A #{len(st.session_state.chat_history) - idx}", expanded=(idx == 0)):
                st.markdown("**Question:**")
                st.write(interaction['question'])
                st.markdown("**Response:**")
                st.write(interaction['response'])

if __name__ == "__main__":
    main()